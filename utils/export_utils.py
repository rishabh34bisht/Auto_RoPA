# utils/export_utils.py
import pandas as pd
from io import BytesIO
from docx import Document

def create_docx(text: str) -> BytesIO:
    """Converts markdown text to a Word Document."""
    doc = Document()
    for line in text.split('\n'):
        if line.startswith('# '):
            doc.add_heading(line.replace('# ', ''), 0)
        elif line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), 1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), 2)
        elif line.startswith('- '):
            doc.add_paragraph(line.replace('- ', ''), style='List Bullet')
        elif line.strip() == '':
            continue
        else:
            doc.add_paragraph(line)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_excel(data_list: list) -> BytesIO:
    """Converts list of dicts to a formatted Excel file."""
    df = pd.DataFrame(data_list)
    buffer = BytesIO()
    
    # Use xlsxwriter engine for formatting
    with pd.ExcelWriter(buffer, engine='xlsxwriter') as writer:
        df.to_excel(writer, sheet_name='ROPA', index=False)
        
        workbook = writer.book
        worksheet = writer.sheets['ROPA']
        
        # Formats
        header_format = workbook.add_format({
            'bold': True,
            'text_wrap': True,
            'valign': 'top',
            'bg_color': '#D9E1F2',
            'border': 1
        })
        
        cell_format = workbook.add_format({
            'text_wrap': True,
            'valign': 'top',
            'border': 1
        })
        
        # Apply formats
        for col_num, value in enumerate(df.columns.values):
            worksheet.write(0, col_num, value, header_format)
            worksheet.set_column(col_num, col_num, 25, cell_format) # Auto-width approximation
            
        worksheet.freeze_panes(1, 0) # Freeze header row
        
    buffer.seek(0)
    return buffer